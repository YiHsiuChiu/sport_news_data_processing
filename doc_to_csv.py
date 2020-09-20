import os #用於獲取目標文件所在路徑
from win32com import client as wc #用於開啟doc 並轉檔成docx
from docx import Document #用於讀取docx檔案
import pandas as pd
import shutil
import sys #用於讀取命令列參數

path=os.getcwd()+"\\data\\105\\10"
#path=sys.argv[1] #資料input路徑
files=[] #原始odt檔案名稱array
docx_files=[] #docx檔案名稱array
for file in os.listdir(path):  #找尋此資料夾內的所有檔案
    if file.endswith(".odt"): #排除資料夾內的其它檔案，只獲取".odt"的word檔案
        files.append(file) 
        docx_files.append(file[:-3]+'docx')      

#刪除docx資料夾
if os.path.exists(path+'/docx'):
    try:
        shutil.rmtree(path+'\docx')
    except OSError as e:
        print(e)
    else:
        print("docx directory is deleted successfully")
    
#刪除csv資料夾
if os.path.exists(path+'/csv'):
    try:
        shutil.rmtree(path+'\csv')
    except OSError as e:
        print(e)
    else:
        print("output directory is deleted successfully")

#建立docx資料夾
os.mkdir(path+'/docx')  
    
word = wc.Dispatch("Word.Application") # 打開word程式
for file in files:
    doc = word.Documents.Open(path+'\\'+file) #打開doc word檔案
    print(path+"\\docx\\"+file[:-3]+'docx'.format(file))
    doc.SaveAs(path+"\\docx\\"+file[:-3]+'docx'.format(file), 12) #另存為後綴為".docx"的檔案，其中參數12指docx檔案
    print(file+"轉檔完成！")
    doc.Close() #關閉原來word檔案
word.Quit()
print("所有轉檔完成！")

#建立output資料夾
os.mkdir(path+'/csv')


#處理所有docx檔案
for docx_file in docx_files:
    document = Document(path+'/docx'+ "\\"+docx_file)
    
    #從title中找到主播、日期等資訊
    needData = ['-1']
    title=document.paragraphs[1].text.split()
    #若資料有空格，則資料在下一段落
    if not '主播' in title[0]:
        title=document.paragraphs[2].text.split()
    #print(title)
    for i in (0,len(title)-1):
        if len(title[i])<2:
            continue
        if '主播' in title[i]:   
            needData.append(title[i])
        if '年' in title[i] and '月' in title[i] and '日' in title[i]:
            needData.append(title[i])
            break
    #使額外資訊符合row格式
    needData.append('');
    needData.append('');
    needData.append('');
    #print(needData)
    
    #所有資料
    totalData = []
    totalData.append(needData) 
    
    columns = [] #每個資料的欄位名稱
    for cell in document.tables[0].rows[0].cells:
        columns.append(cell.text) 
    #print(columns)
    #print(document.tables[0].rows[5].cells[3].text)
    for table in document.tables: #將docx檔案裡的table取出
        for row in table.rows:
            rowData = [] #每行資料
            for cell in row.cells:
                rowData.append(cell.text.replace("\'", "’").replace("\"", "”")) #每行裡面的每個資料，並去掉符號("、')
            if rowData == columns:
                continue
            totalData.append(rowData)
    #print(totalData)
    
    for i in range(0,len(totalData)):
        for j in range(0,len(totalData[i])):
            totalData[i][j] = ''.join(totalData[i][j].split())
    for i in range(0,len(columns)):
        columns[i] = ''.join(columns[i].split())

    outputDF = pd.DataFrame(data=totalData,columns=columns) #儲存成dataframe格式 以便之後的數據分析 
    #print(outputDF)
    
    
    #暫時先處儲存成csv檔 之後再讀取後好處理
    outputDF.to_csv(path+'/csv/'+ docx_file[:-5] + '.csv', encoding='utf_8_sig')
    print( docx_file[:-5] + '.csv'+"儲存成csv檔完成")
